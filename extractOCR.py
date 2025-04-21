import fitz  # PyMuPDF
import pytesseract
from pytesseract import Output
from PIL import Image
import io
from docx import Document

def extract_text_and_images_from_pdf(pdf_path, docx_path, header_y_limit=100, footer_y_limit=750):
    doc = fitz.open(pdf_path)
    docx_document = Document()

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        docx_document.add_paragraph(f"Page {page_num + 1}")

        # Extracting text and filtering out headers/footers
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    bbox = span["bbox"]
                    y0, y1 = bbox[1], bbox[3]
                    text = span["text"]

                    # Skip headers and footers based on y-coordinate
                    if y0 > header_y_limit and y1 < footer_y_limit:
                        docx_document.add_paragraph(text)

        # Process images
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            # Convert image to text using pytesseract
            image = Image.open(io.BytesIO(image_bytes))
            text_from_image = pytesseract.image_to_string(image, lang="spa", config='--psm 6')

            # Filter out header/footer text in images if needed (optional)
            docx_document.add_paragraph(f"Image {img_index + 1} OCR Text:")
            docx_document.add_paragraph(text_from_image)

            # Optionally, save the image or embed it in the docx
            image_filename = f"image_{page_num + 1}_{img_index + 1}.png"
            image.save(image_filename)
            docx_document.add_paragraph(f"[Image saved as {image_filename}]")

    docx_document.save(docx_path)
    print(f"PDF text and images have been saved to {docx_path}")

# Example usage
pdf_path = "input.pdf"       # Path to the PDF file
docx_path = "output.docx"    # Path to save the .docx file
extract_text_and_images_from_pdf(pdf_path, docx_path)
